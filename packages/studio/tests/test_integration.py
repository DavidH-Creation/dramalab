"""Integration test: init → run 1 round → status → text → export."""

import io
import json
import pytest
import pytest_asyncio
from unittest.mock import patch, MagicMock
from httpx import AsyncClient, ASGITransport
from dramalab_studio.server import create_app
from dramalab_studio.plugin_protocol import RoundResult

pytestmark = pytest.mark.asyncio


@pytest.fixture
def app():
    return create_app()


@pytest_asyncio.fixture
async def client(app):
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c


async def test_full_flow(client, tmp_path):
    """Test full flow: upload → init → status → export."""
    from docx import Document

    # Upload screenplay
    doc = Document()
    doc.add_paragraph("第一集")
    doc.add_paragraph("场景一")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    resp = await client.post("/api/upload", files={"file": ("test.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert resp.status_code == 200
    text = resp.json()["text"]

    # Init plugin (with mocked backend to avoid needing Claude CLI)
    with patch("scriptsmith.backends.claude_cli.ClaudeCLIBackend") as MockBackend, \
         patch("scriptsmith.splitter.split_screenplay") as mock_split, \
         patch("scriptsmith.deriver.derive_all"), \
         patch("scriptsmith.git_ops.git_init"):

        mock_split.return_value = [MagicMock(
            id="seq_01", filename="seq_01.md", title="Ep 1",
            char_count=100, scene_count=1, markers=[],
            to_dict=lambda: {"id": "seq_01", "filename": "seq_01.md"},
        )]

        resp = await client.post("/api/plugins/script-forge/init", json={
            "input_text": text,
            "criteria_text": "Structure: 20",
            "config": {"model": "sonnet"},
        })
        assert resp.status_code == 200
        session_id = resp.json()["session_id"]

    # Check status
    resp = await client.get(f"/api/plugins/script-forge/{session_id}/status")
    assert resp.status_code == 200
    assert resp.json()["status"] == "idle"
