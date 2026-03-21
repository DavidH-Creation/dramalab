import io
import pytest
from docx import Document

pytestmark = pytest.mark.asyncio

async def test_upload_docx(client):
    """Upload a .docx file and get extracted text."""
    doc = Document()
    doc.add_paragraph("第一集 死神来了")
    doc.add_paragraph("场景一：殡仪馆")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    resp = await client.post(
        "/api/upload",
        files={"file": ("test.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200
    data = resp.json()
    assert "第一集 死神来了" in data["text"]
    assert data["filename"] == "test.docx"

async def test_upload_md(client):
    """Upload a .md file and get text as-is."""
    content = b"# Scoring Criteria\n\nStructure: 20 points"
    resp = await client.post(
        "/api/upload",
        files={"file": ("criteria.md", io.BytesIO(content), "text/markdown")},
    )
    assert resp.status_code == 200
    assert "Scoring Criteria" in resp.json()["text"]

async def test_upload_invalid_type(client):
    """Reject unsupported file types."""
    resp = await client.post(
        "/api/upload",
        files={"file": ("image.png", io.BytesIO(b"fake"), "image/png")},
    )
    assert resp.status_code == 400
