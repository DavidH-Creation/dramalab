# tests/test_exporter.py
"""Tests for script_forge.exporter."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from script_forge.exporter import export_to_docx


class TestExportToDocx:
    def test_creates_docx(self, tmp_workspace: Path):
        output = tmp_workspace / "exports" / "output.docx"
        export_to_docx(tmp_workspace, output)
        assert output.exists()

    def test_contains_sequence_text(self, tmp_workspace: Path):
        output = tmp_workspace / "exports" / "output.docx"
        export_to_docx(tmp_workspace, output)

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "张陵醒来" in full_text

    def test_sequences_in_order(self, tmp_workspace: Path):
        output = tmp_workspace / "exports" / "output.docx"
        export_to_docx(tmp_workspace, output)

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        pos1 = full_text.index("张陵醒来")
        pos2 = full_text.index("新的冒险")
        assert pos1 < pos2
