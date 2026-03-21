# src/script_forge/exporter.py
"""Export sequences back to a single docx file."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from script_forge.workspace import load_manifest


def export_to_docx(workspace: Path, output_path: Path) -> None:
    """Merge all sequences into a single docx file.

    Text-level round-trip only — does not preserve original docx styling.
    """
    sequences = load_manifest(workspace)
    doc = Document()

    for seq in sequences:
        seq_path = workspace / "sequences" / seq.filename
        text = seq_path.read_text(encoding="utf-8")

        # Add sequence title as heading
        doc.add_heading(seq.title, level=2)

        # Add paragraphs
        for paragraph in text.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
